import os
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from xml.etree import ElementTree as ET
import zipfile
from PIL import Image
from docx.shared import Inches
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING


def process_image_diagram_formatter(source_doc):
    # source_doc = "source — копия — копия.docx"
    # source_doc = "КП_Чумаров_Черновик (10_12_2024)).docx"
    source_document = Document(source_doc)
    source_zip = zipfile.ZipFile(source_doc, 'r')

    images = [file for file in source_zip.namelist() if file.startswith('word/media/')]
    print("Порядок изображений до сортировки:", images)
    images = sorted(
        [file for file in source_zip.namelist() if file.startswith('word/media/')], 
        key=lambda x: int(''.join(filter(str.isdigit, x))) if ''.join(filter(str.isdigit, x)).isdigit() else 0
    )
    print("Порядок изображений после сортировки:", images)
    images_index = 0 
    num_chart = 0
    image_chart_num = 0


    def get_paragraph_xml(paragraph):
        return paragraph._element.xml

    def is_paragraph_image(paragraph):
        return bool(paragraph._element.xpath('.//pic:pic'))

    def get_image_dimensions_from_word(paragraph_xml):
        try:
            root = ET.fromstring(paragraph_xml)
            drawing = root.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing')
            if drawing is not None:
                extent = drawing.find('.//{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}extent')
                if extent is not None:
                    cx = int(extent.attrib['cx']) if 'cx' in extent.attrib else 0
                    cy = int(extent.attrib['cy']) if 'cy' in extent.attrib else 0
                    pixels_per_mm = 96 / 25.4
                    width_px = int(cx * pixels_per_mm / 100)
                    height_px = int(cy * pixels_per_mm / 100)
                    return width_px, height_px
        except Exception as e:
            print(f"Ошибка при извлечении размеров из XML: {e}")
        return None, None

    def is_paragraph_chart(paragraph_xml):
        try:
            root = ET.fromstring(paragraph_xml)
            return root.find('.//{http://schemas.openxmlformats.org/drawingml/2006/chart}chart') is not None
        except ET.ParseError:
            return False

    def center_chart(paragraph):
        # Устанавливаем выравнивание для параграфа
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Получаем XML элемента параграфа
        p_element = paragraph._element
        drawing_element = p_element.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing')
        if drawing_element is not None:
            # Ищем <wp:anchor> или <wp:inline>
            anchor_or_inline = drawing_element.find('.//{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}anchor') or \
                            drawing_element.find('.//{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}inline')
            if anchor_or_inline is not None:
                # Добавляем/изменяем элемент <wp:positionH>
                position_h = anchor_or_inline.find('{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}positionH')
                if position_h is None:
                    position_h = OxmlElement('wp:positionH')
                    position_h.set('relativeFrom', 'page')
                    anchor_or_inline.insert(0, position_h)
                align = position_h.find('{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}align')
                if align is None:
                    align = OxmlElement('wp:align')
                    position_h.append(align)
                align.text = 'center'

                # Обновляем атрибуты привязки графика
                anchor_or_inline.set('allowOverlap', '1')
                anchor_or_inline.set('relativeHeight', '0')
                print("График отцентрирован в объекте")

    def resize_images_in_doc(doc, max_width):
        """
        Изменяет размеры изображений в Word-документе, чтобы их ширина не превышала заданное значение.
        :param doc: Объект документа (Document).
        :param max_width: Максимальная ширина изображения в дюймах.
        """
        max_width_emus = int(max_width * 914400)  # Конвертация дюймов в EMU (единицы Word)
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                # Проверяем, содержит ли run изображение
                if run._element.xpath('.//pic:pic'):
                    inline_shapes = run._element.xpath('.//wp:extent')
                    for extent in inline_shapes:
                        # Получение текущей ширины и высоты
                        width = int(extent.get('cx'))
                        height = int(extent.get('cy'))

                        if width > max_width_emus:
                            # Уменьшаем ширину до max_width, сохраняя пропорции
                            scaling_factor = max_width_emus / width
                            new_width = max_width_emus
                            new_height = int(height * scaling_factor)

                            # Устанавливаем новые размеры
                            extent.set('cx', str(new_width))
                            extent.set('cy', str(new_height))


    prev_para_xml = None
    for para in source_document.paragraphs:
        para_xml = get_paragraph_xml(para)
        if prev_para_xml:
            if para.text.strip():
                if prev_para_xml_num == 1:
                # if is_paragraph_image(prev_para_xml) and para.text.strip():
                    prev_para_xml = None
                    prev_para_xml_num = 0
                    current_text = para.text.strip().capitalize()
                    para.clear()
                    run = para.add_run(f"Рисунок {image_chart_num} – {current_text}")
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(14)
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER

                    parent = para._element.getparent()
                    current_index = parent.index(para._element)
                    next_element = parent[current_index + 1]
                    if not next_element.xpath('.//w:t') or all(t.text.strip() == "" for t in next_element.xpath('.//w:t')):
                        print("Следующий параграф пуст, новый не создаем.")
                    else:
                        print("Следующий параграф не пуст, создаем новый пустой параграф.")
                        new_para = OxmlElement("w:p")
                        parent.insert(current_index + 1, new_para)

                if prev_para_xml_num == 2:
                # if is_paragraph_chart(prev_para_xml):
                    prev_para_xml = None
                    prev_para_xml_num = 0
                    current_text = para.text.strip().capitalize()
                    para.clear()
                    run = para.add_run(f"Рисунок {image_chart_num} – {current_text}")
                    # num_chart+=1
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(14)
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER

                    parent = para._element.getparent()
                    current_index = parent.index(para._element)
                    next_element = parent[current_index + 1]
                    if not next_element.xpath('.//w:t') or all(t.text.strip() == "" for t in next_element.xpath('.//w:t')):
                        print("Следующий параграф пуст, новый не создаем.")
                    else:
                        print("Следующий параграф не пуст, создаем новый пустой параграф.")
                        new_para = OxmlElement("w:p")
                        parent.insert(current_index + 1, new_para)

        if is_paragraph_image(para) and not is_paragraph_chart(para_xml):
            print("Параграф содержит изображение")

            parent = para._element.getparent()
            current_index = parent.index(para._element)
            previous_element = parent[current_index - 1]
            if not previous_element.xpath('.//w:t') or all(t.text.strip() == "" for t in previous_element.xpath('.//w:t')):
                print("Предыдущий параграф пуст, новый не создаем.")
            else:
                print("Предыдущий параграф не пуст, создаем новый пустой параграф.")
                new_para = OxmlElement("w:p")
                parent.insert(current_index, new_para)

            para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            para.paragraph_format.space_after = Pt(0)

            if images_index < len(images):
                image_chart_num += 1
                print(f"{image_chart_num - 1} image_chart_num (вместо images_index)")
                width, height = get_image_dimensions_from_word(para_xml)
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                width, height = get_image_dimensions_from_word(para_xml)
                if width and height:
                    print(f"Изображение {image_chart_num - 1}: ширина {width}px, высота {height}px")
                else:
                    print("Не удалось извлечь размеры изображения")
            else:
                print("Изображение не найдено в архиве.")
        if is_paragraph_chart(para_xml):
            print("Параграф содержит график")
            image_chart_num += 1
            parent = para._element.getparent()
            current_index = parent.index(para._element)
            previous_element = parent[current_index - 1]
            if not previous_element.xpath('.//w:t') or all(t.text.strip() == "" for t in previous_element.xpath('.//w:t')):
                print("Предыдущий параграф пуст, новый не создаем.")
            else:
                print("Предыдущий параграф не пуст, создаем новый пустой параграф.")
                new_para = OxmlElement("w:p")
                parent.insert(current_index, new_para)

            para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            para.paragraph_format.space_after = Pt(0)

            center_chart(para)
        # prev_para_xml = para if is_paragraph_image(para) and not is_paragraph_chart(para_xml) else para_xml if is_paragraph_chart(para_xml) and not is_paragraph_image(para) else None
        if is_paragraph_image(para) and not is_paragraph_chart(para_xml):
            prev_para_xml = para
            prev_para_xml_num = 1
        elif is_paragraph_chart(para_xml):
            prev_para_xml = para_xml
            prev_para_xml_num = 2
        else:
            prev_para_xml = None
            prev_para_xml_num = 0

    # Получаем ширину страницы
    page_width = source_document.sections[0].page_width.inches - source_document.sections[0].left_margin.inches - source_document.sections[0].right_margin.inches

    # Изменяем размеры изображений
    resize_images_in_doc(source_document, max_width=page_width)

    # Сохраняем изменения
    source_document.save(source_doc)
    print("Готово")
