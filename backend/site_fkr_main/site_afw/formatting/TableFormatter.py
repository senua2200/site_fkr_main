from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_PARAGRAPH_ALIGNMENT
from docx.table import Table


def process_table_formatter(source_doc):
    # source_doc = "КП_Чумаров_Черновик (10_12_2024)).docx"
    source_document = Document(source_doc)

    # Список исключенных заголовков
    excluded_titles = {
        "введение", "Список литературы", "Приложение", "список использованной литературы", "Список используемой литературы", "Литература", 
        "Библиографический список", "Перечень источников", "Источники и литература", "Список источников и литературы", "Рекомендуемая литература", 
        "Привлеченные источники", "Список использованных источников и материалов", "Список использованных источников и литературы", "Список использованных источников",  
        "Использованные источники", "Используемые источники", "Список использованных источников"
    }

    def is_next_element_table(paragraph):
        """
        Проверяет, является ли следующий элемент после параграфа таблицей.
        """
        xml_element = paragraph._element
        next_element = xml_element.getnext()
        if next_element is not None and len(next_element.tag) > 0 and next_element.tag == qn('w:tbl'):
            parent = next_element.getparent()
            next_index = parent.index(next_element)
            element_behind_table = parent[next_index + 1]
            if not element_behind_table.xpath('.//w:t') or all(t.text.strip() == "" for t in element_behind_table.xpath('.//w:t')):
                print("Следующий параграф пуст, новый не создаем.")
            else:
                print("Следующий параграф не пуст, создаем новый пустой параграф.")
                # Создаем новый пустой параграф с нужными настройками
                new_para = OxmlElement("w:p")
                new_para_pPr = OxmlElement("w:pPr")

                # Устанавливаем параметры форматирования
                spacing = OxmlElement("w:spacing")
                spacing.set(qn("w:line"), "360")  # Одинарный интервал
                spacing.set(qn("w:before"), "0")  # 0 пт до
                spacing.set(qn("w:after"), "0")   # 0 пт после
                new_para_pPr.append(spacing)

                # Добавляем параметры в новый параграф
                new_para.append(new_para_pPr)

                # Вставляем новый параграф после текущего элемента
                parent.insert(next_index + 1, new_para)
                
            apply_table_style(next_element)

            return True
        return False

    def apply_table_style(table):
        """
        Применяет стиль ко всем ячейкам таблицы: шрифт 'Times New Roman', размер 12, одинарный интервал, 0 пт до и после.
        """
        table = Table(table, source_document)
        dash = False
        for row in table.rows:
            for cell in row.cells:
                if not cell.text.strip():
                    cell.text = "–"
                    dash = True
                for paragraph in cell.paragraphs:
                    # Устанавливаем шрифт и размер
                    if dash:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        dash = False

                    for run in paragraph.runs:
                        run.font.name = 'Times New Roman'  # Шрифт
                        run.font.size = Pt(12)  # Размер шрифта

                    # Устанавливаем интервал и отступы
                    paragraph.paragraph_format.line_spacing = 1  # Одинарный интервал
                    paragraph.paragraph_format.space_before = Pt(0)  # 0 пт до
                    paragraph.paragraph_format.space_after = Pt(0)  # 0 пт после

    def is_highlighted_yellow(paragraph):
        """
        Проверяет, выделен ли параграф желтым цветом.
        """
        highlight_elements = paragraph._element.xpath('.//w:highlight')
        for elem in highlight_elements:
            if elem.get(qn('w:val')) == 'yellow':  # Проверяем значение выделения
                return True
        return False

    # Обработка документа
    table_number = 0
    processing = False  # Флаг для управления диапазоном обработки

    for para in source_document.paragraphs:
        # Проверка на желтое выделение
        if is_highlighted_yellow(para):
            text = para.text.strip()
            if text.lower() in (title.lower() for title in excluded_titles):
                print(f"Обнаружен заголовок: {text}")
                # Переключаем флаг
                if not processing:
                    print("Начинаем обработку.")
                    processing = True
                else:
                    print("Заканчиваем обработку.")
                    processing = False
                    break  # Завершаем цикл, если больше не нужно искать таблицы

        # Если находимся в режиме обработки, проверяем наличие таблиц
        if processing and is_next_element_table(para):
            print("ТАБЛИЦА НАЙДЕНА")
            table_number += 1
            if para.text.strip():
                print('Текст есть')

                # Проверка предыдущего параграфа на пустоту
                parent = para._element.getparent()
                current_index = parent.index(para._element)
                previous_element = parent[current_index - 1]
                if not previous_element.xpath('.//w:t') or all(t.text.strip() == "" for t in previous_element.xpath('.//w:t')):
                    print("Предыдущий параграф пуст, новый не создаем.")
                else:
                    print("Предыдущий параграф не пуст, создаем новый пустой параграф.")
                    # Создаем новый пустой параграф с нужными настройками
                    new_para = OxmlElement("w:p")
                    new_para_pPr = OxmlElement("w:pPr")

                    # Устанавливаем параметры форматирования
                    spacing = OxmlElement("w:spacing")
                    spacing.set(qn("w:line"), "360")  # Одинарный интервал
                    spacing.set(qn("w:before"), "0")  # 0 пт до
                    spacing.set(qn("w:after"), "0")   # 0 пт после
                    new_para_pPr.append(spacing)

                    # Добавляем параметры в новый параграф
                    new_para.append(new_para_pPr)

                    # Вставляем новый параграф после текущего элемента
                    parent.insert(current_index, new_para)

                para.text = f"Таблица – {table_number} {para.text[0].upper() + para.text[1:]}" 
                para.paragraph_format.line_spacing = 1.5
                para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                for run in para.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(14)
    print("ТАБЛИЦЫ!!!")
    source_document.save(source_doc)
