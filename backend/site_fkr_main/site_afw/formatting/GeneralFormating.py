from docx import Document
from docx.shared import Pt, Cm, RGBColor, Mm
from docx.enum.text import WD_LINE_SPACING, WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import zipfile
import os
import shutil
import win32com.client as win32
from lxml import etree
from docx.enum.text import WD_BREAK  # <-- Добавляем импорт WD_BREAK

def process_GeneralFormating(file_path, hyperlink_mode = 1):
    #----------------------НАСТРОЙКИ И ПАРАМЕТРЫ--------------------------
    # 0 – удалять гиперссылки, 1 – применять стиль
    hyperlink = 1

    # Параметры форматирования
    font_size = 14                   # размер основного текста
    font_size_headlines = 16         # размер заголовков
    font_name = 'Times New Roman'
    para_alignment = 3               # выравнивание по ширине (полное)
    para_alignment_headlines = 1     # выравнивание по центру для заголовков
    first_line_indent = 1.25         # отступ первой строки в см
    line_spacing_rule_1_5 = WD_LINE_SPACING.SINGLE # (WD_LINE_SPACING.ONE_POINT_FIVE  межстрочный интервал 1.5) (WD_LINE_SPACING.SINGLE междусточный одинарный интервал)

    # Параметры полей (в мм)
    topSection = 20
    bottomSection = 20
    leftSection = 30
    rigthSection = 10

    # Отступы (в пунктах)
    left_margin = 0 
    Right_margin = 0
    interval_before_paragraph = 0
    interval_after_paragraph = 0

    # Название общего стиля текста (убедитесь, что стиль доступен в документе)
    style_text = "Normal" 

    #----------------------ФУНКЦИИ ДЛЯ ФОРМАТИРОВАНИЯ--------------------------
    def set_paragraph_style(para, para_alignment, first_line_indent, line_spacing, font_size, font_name):
        """
        Применяет стиль к параграфу и форматирует текст.
        Если в абзаце содержится изображение, отступ первой строки не задаётся.
        """
        para.style = style_text
        para.alignment = para_alignment
        # 🚫 Пропускаем абзацы с изображениями
        if para._p.xpath('.//w:drawing'):
            print(f"Пропущен абзац с изображением: \"{para.text[:30]}...\"")
            return

        para.paragraph_format.line_spacing_rule = line_spacing
        para.paragraph_format.left_indent = Pt(left_margin)
        para.paragraph_format.right_indent = Pt(Right_margin)
        para.paragraph_format.space_before = Pt(interval_before_paragraph)
        para.paragraph_format.space_after = Pt(interval_after_paragraph)
        
        for run in para.runs:
            run.font.size = Pt(font_size)
            run.font.name = font_name
            run.font.color.rgb = RGBColor(0, 0, 0)
            run.bold = False
            run.italic = False
            run.underline = False

            if run.font.highlight_color == 7:  # Проверяем, выделен ли текст желтым цветом
                    # Добавляем разрыв страницы перед текущим абзацем
                    prior_paragraph = para.insert_paragraph_before()
                    prior_paragraph.add_run().add_break(WD_BREAK.PAGE)  # <-- Исправлено
                    para.paragraph_format.first_line_indent = Cm(0)
                    print(f"ЭТО УЖЕ ПРОВЕРКА НА ВЫДЕЛЕНИЕ Текст: {run.text}, Цвет выделения: {run.font.highlight_color}")
                    break  # Достаточно одного разрыва перед абзацем

    def apply_styles_to_paragraphs_after_keyword(doc, keyword):
        """
        Применяет указанный стиль к параграфам документа, начиная с абзаца,
        где найдено ключевое слово (сравнение без учёта регистра).
        Если ключевое слово не найдено, стиль применяется ко всему документу.
        """
        section = doc.sections[0]
        section.left_margin = Mm(leftSection)
        section.right_margin = Mm(rigthSection)
        section.top_margin = Mm(topSection)
        section.bottom_margin = Mm(bottomSection)

        keyword_found = False

        for para in doc.paragraphs:

            if para._p.xpath('.//w:drawing'):
                print(f"Пропущен абзац с изображением: \"{para.text[:30]}...\"")
                continue

            if keyword_found:
                set_paragraph_style(
                    para,
                    para_alignment=para_alignment,
                    first_line_indent=first_line_indent,
                    line_spacing=line_spacing_rule_1_5,
                    font_size=font_size,
                    font_name=font_name
                )
            elif para.text.strip().lower() == keyword.lower():
                keyword_found = True
                set_paragraph_style(
                    para,
                    para_alignment=para_alignment_headlines,
                    first_line_indent=first_line_indent,
                    line_spacing=line_spacing_rule_1_5,
                    font_size=20,  # можно заменить на font_size_headlines
                    font_name=font_name
                )
        
        if not keyword_found:
            for para in doc.paragraphs:
                
                if para._p.xpath('.//w:drawing'):
                    print(f"Пропущен абзац с изображением: \"{para.text[:30]}...\"")
                    continue

                set_paragraph_style(
                    para,
                    para_alignment=para_alignment,
                    first_line_indent=first_line_indent,
                    line_spacing=line_spacing_rule_1_5,
                    font_size=font_size,
                    font_name=font_name
                )

    #----------------------ФУНКЦИИ ДЛЯ РАБОТЫ С ГИПЕРССЫЛКАМИ--------------------------
    def update_hyperlink_style(paragraph):
        """
        Применяет стиль гиперссылок к параграфу:
        - Шрифт: Times New Roman
        - Размер: 14 пунктов (значение в половинках пункта – 28)
        - Цвет: синий (0000FF)
        - Подчеркивание: одинарное
        """
        for hyperlink in paragraph._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}hyperlink'):
            for run in hyperlink.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}r'):
                rPr = run.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rPr')
                if rPr is None:
                    rPr = OxmlElement('w:rPr')
                    run.insert(0, rPr)

                rFonts = OxmlElement('w:rFonts')
                rFonts.set(qn('w:ascii'), 'Times New Roman')
                rFonts.set(qn('w:eastAsia'), 'Times New Roman')
                rFonts.set(qn('w:hAnsi'), 'Times New Roman')
                rPr.append(rFonts)

                size = OxmlElement('w:sz')
                size.set(qn('w:val'), '28')
                rPr.append(size)

                color = OxmlElement('w:color')
                color.set(qn('w:val'), "0000FF")
                rPr.append(color)

                underline = OxmlElement('w:u')
                underline.set(qn('w:val'), "single")
                rPr.append(underline)

    def remove_hyperlinks_keep_text_inplace(file_path, start_index=None):
        """
        Удаляет гиперссылки из документа, оставляя только текст.
        Если указан start_index, обработка ведётся начиная с этого параграфа.
        """
        temp_dir = 'temp_docx'
        with zipfile.ZipFile(file_path, 'r') as zip_ref:
            zip_ref.extractall(temp_dir)

        document_xml_path = os.path.join(temp_dir, 'word', 'document.xml')
        tree = etree.parse(document_xml_path)
        ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}

        if start_index is not None:
            paragraphs = tree.xpath('//w:body//w:p[position()>={} and not(ancestor::w:hyperlink)]'.format(start_index + 1), namespaces=ns)
        else:
            paragraphs = tree.xpath('//w:body//w:p[not(ancestor::w:hyperlink)]', namespaces=ns)

        for para in paragraphs:
            for hyperlink in para.xpath('.//w:hyperlink', namespaces=ns):
                parent = hyperlink.getparent()
                for child in list(hyperlink):
                    parent.insert(parent.index(hyperlink), child)
                parent.remove(hyperlink)

        with open(document_xml_path, 'wb') as xml_file:
            tree.write(xml_file, xml_declaration=True, encoding='utf-8')

        with zipfile.ZipFile(file_path, 'w') as zip_ref:
            for foldername, subfolders, filenames in os.walk(temp_dir):
                for filename in filenames:
                    full_path = os.path.join(foldername, filename)
                    arcname = os.path.relpath(full_path, temp_dir)
                    zip_ref.write(full_path, arcname)

        shutil.rmtree(temp_dir)
        print("Гиперссылки успешно удалены.")

    def process_document(docx_path, hyperlink):
        """
        Обрабатывает документ DOCX:
        - Если параметр hyperlink == 0, гиперссылки удаляются (оставляя текст).
        - Если параметр hyperlink == 1, к гиперссылкам применяется заданный стиль.
        Обработка начинается после нахождения ключевого слова "Введение", выделенного жёлтым.
        Если ключевое слово не найдено, обрабатываются все гиперссылки в документе.
        """
        document = Document(docx_path)
        start_index = None
        for i, para in enumerate(document.paragraphs):
            for run in para.runs:
                if "Введение" in run.text and run.font.highlight_color == 7:
                    start_index = i
                    print(f'Найдено выделенное жёлтым "Введение" в параграфе {i}: "{para.text}"')
                    break
            if start_index is not None:
                break

        if hyperlink == 0:
            if start_index is not None:
                remove_hyperlinks_keep_text_inplace(docx_path, start_index)
            else:
                print('Ключевое слово "Введение" не найдено. Удаляем все гиперссылки в документе.')
                remove_hyperlinks_keep_text_inplace(docx_path)
        elif hyperlink == 1:
            if start_index is not None:
                relevant_paragraphs = document.paragraphs[start_index+1:]
                print(f"Применяем стиль к гиперссылкам после параграфа {start_index}...")
            else:
                print('Ключевое слово "Введение" не найдено. Применяем стиль ко всем гиперссылкам.')
                relevant_paragraphs = document.paragraphs

            for para in relevant_paragraphs:
                update_hyperlink_style(para)

            document.save(docx_path)
            print("Стиль гиперссылок успешно применён.")
        else:
            print("Ошибка: неверно указан параметр для работы с гиперссылками. Используйте 0 для удаления или 1 для применения стиля.")

    #----------------------ОСНОВНАЯ ЧАСТЬ ПРОГРАММЫ--------------------------
    print('Началась работа GeneralFormaning')
    # 1. Обработка гиперссылок (в зависимости от значения переменной hyperlink)
    docx_file = file_path  # Путь к документу
    doc = Document(docx_file)
    process_document(docx_file, hyperlink_mode)

    # 2. Применяем общее форматирование текста, начиная с ключевого слова "Введение"
    docx_file = file_path
    doc = Document(docx_file)
    apply_styles_to_paragraphs_after_keyword(doc, "Введение")
    doc.save(docx_file)
    print("Форматирование завершено.")
    print('Работа GeneralFormaning завершилась')


