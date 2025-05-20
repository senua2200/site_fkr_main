import win32com.client
from docx import Document
import os
from docx.enum.text import WD_BREAK
import win32com.client

def process_page_numbering(file_path,numlist):
    def clear_footer_and_numbering(doc_path):
        doc = Document(doc_path)
        
        for section in doc.sections:
            section.footer.is_linked_to_previous = False
            footer = section.footer
            if footer:
                for paragraph in footer.paragraphs:
                    for run in paragraph.runs:
                        run.clear()
                while footer.paragraphs:
                    p = footer.paragraphs[0]
                    p._element.getparent().remove(p._element)
                footer.add_paragraph()
        
        doc.save(doc_path)
        print("Очистка колонтитулов завершена")

    def remove_section_breaks(name_doc_os):
        try:
            current_dir = os.getcwd()
            input_file_path = os.path.join(current_dir, name_doc_os)

            word = win32com.client.Dispatch("Word.Application")
            word.Visible = False

            doc = word.Documents.Open(input_file_path)
            doc.Activate()

            find = doc.Content.Find
            find.ClearFormatting()
            find.Replacement.ClearFormatting()
            find.Execute(FindText="^b", ReplaceWith="", Replace=2)

            doc.SaveAs(input_file_path)
            doc.Close(False)
            word.Quit()

            print("Удалены разрывы разделов")
        except Exception as e:
            print(f"Произошла ошибка: {e}")

    def find_first_heading1_win32(file_path):
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        doc = word.Documents.Open(file_path)

        for i in range(1, doc.Paragraphs.Count + 1):
            para = doc.Paragraphs(i)
            if para.Style.NameLocal == "Заголовок 1":  # Проверяем выделение желтым цветом
                page_num = para.Range.Information(3)  # Получаем номер страницы
                doc.Close(False)
                word.Quit()
                return page_num
        
        doc.Close(False)
        word.Quit()
        return None

    def add_page_numbers(file_name, start_page):
        file_path = os.path.join(os.getcwd(), file_name)
        
        if start_page == 0:
            start_page = find_first_heading1_win32(file_path) or 1
        
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        
        doc = word.Documents.Open(file_path)
        selection = doc.ActiveWindow.Selection
        selection.GoTo(1, 1, start_page)

        if selection.Information(8):
            selection.MoveDown(Unit=5, Count=1)
        
        if selection.Paragraphs.Count > 0 and selection.Paragraphs.First.Range.Style:
            if selection.Paragraphs.First.Range.Style.NameLocal.startswith("Заголовок"):
                selection.MoveDown(Unit=5, Count=1)
                print('Заголовок найден, смещение курсора')

        selection.Collapse(0)
        
        if start_page > 1:
            try:
                selection.InsertBreak(3)
            except Exception as e:
                print("Ошибка вставки разрыва:", e)
            section_index = selection.Information(2)
        else:
            section_index = 1
        
        section = doc.Sections(section_index)
        footer = section.Footers(1)
        footer.LinkToPrevious = False
        
        if footer.Range:
            footer.Range.Text = " "
            footer.Range.Fields.Add(footer.Range, win32com.client.constants.wdFieldPage)
            footer.Range.ParagraphFormat.Alignment = win32com.client.constants.wdAlignParagraphCenter
        
        section.Footers(1).PageNumbers.RestartNumberingAtSection = True
        section.Footers(1).PageNumbers.StartingNumber = start_page
        
        doc.Save()
        doc.Close()
        word.Quit()
        print(f'Нумерация добавлена корректно, начиная с {start_page} страницы')

    #word = win32com.client.Dispatch("Word.Application")
    #print(dir(word))  # Выведет список доступных атрибутов
    #print(hasattr(word, "Constants"))  # Проверка наличия констант
    import win32api
    print(win32api.GetVersion())

    name_doc = file_path
    clear_footer_and_numbering(name_doc)
    remove_section_breaks(name_doc)
    add_page_numbers(name_doc, numlist) # последнее число указывает на то с какой страницы необходимо начать нумерацию, если оно равно 0 то нумерация будет с первого желтого цвета выделения текста. т.е с заголовка